import streamlit as st
import os
from openai import OpenAI
from vector_store import create_vector_store
from docx import Document
from docx.shared import Pt
import random

# 🔥 SELECT S "SPECIFICKÉ"
def select_with_custom(label, options):
    choice = st.selectbox(
        label,
        options + ["Specifické"],
        index=None,
        placeholder="Vyber...",
        key=label
    )

    if choice == "Specifické":
        return st.text_input(f"{label} - upřesnění", key=label+"_custom")

    return choice

client = OpenAI(api_key=os.getenv("OPENAI_API_KEY"))

@st.cache_resource
def get_vectorstore():
    return create_vector_store()

vectorstore = get_vectorstore()

st.title("AI Asistent projektanta")

# 🔹 SIDEBAR
data = {}

with st.sidebar:
    st.header("Parametry projektu")

    if st.button("🔄 Reset"):
        st.session_state.clear()
        st.rerun()

    if st.button("🎲 Náhodně vyplnit"):

        st.session_state["Druh stavby"] = random.choice([
            "Rodinný dům", "Dvojdům", "Řadový dům"
        ])

        st.session_state["Lokalita"] = random.choice([
            "Brno", "Praha", "Ostrava", "Plzeň"
        ])

        st.session_state["Typ území"] = random.choice([
            "v zastavěné části", "v nezastavěné části"
        ])

        st.session_state["Počet NP"] = random.choice(["1", "2"])

        st.session_state["Půdorys"] = random.choice([
            "obdélník", "L"
        ])

        st.session_state["Konstrukce"] = random.choice([
            "zděná", "dřevostavba"
        ])

        st.session_state["Vytápění"] = random.choice([
            "TČ vzduch-voda", "plyn"
        ])

        st.session_state["Voda"] = "vodovod"
        st.session_state["Kanalizace"] = "kanalizace"
        st.session_state["Elektřina"] = "nová přípojka"

        st.rerun()

    with st.expander("I. Identita a pozemek", expanded=True):
        data["Druh stavby"] = select_with_custom("Druh stavby", [
            "Rodinný dům", "Dvojdům", "Řadový dům", "Rekreační objekt"
        ])
        data["Lokalita"] = st.text_input("Lokalita", placeholder="např. Brno")
        data["Typ území"] = select_with_custom("Typ území", [
            "v zastavěné části", "v nezastavěné části", "chatová oblast"
        ])
        data["Doprava"] = select_with_custom("Napojení na dopravu", [
            "stávající sjezd", "nový sjezd", "jiný pozemek"
        ])

    with st.expander("II. Tvar a hmota"):
        data["Půdorys"] = select_with_custom("Půdorys", [
            "obdélník", "čtverec", "L", "U"
        ])
        data["Počet NP"] = select_with_custom("Počet NP", ["1", "2", "3"])
        data["Podkroví"] = st.radio("Podkroví", ["ano", "ne"])
        data["Podsklepení"] = st.radio("Podsklepení", ["ne", "ano", "částečně"])
        data["Terén"] = select_with_custom("Terén", [
            "rovina", "mírný sklon", "prudký svah"
        ])
        data["Garáž"] = select_with_custom("Garáž", [
            "není", "součást domu", "samostatná"
        ])

    with st.expander("III. Konstrukce a zdroje"):
        data["Konstrukce"] = select_with_custom("Konstrukce", [
            "zděná", "dřevostavba", "beton", "ocel"
        ])
        data["Vytápění"] = select_with_custom("Vytápění", [
            "TČ vzduch-voda", "plyn", "elektřina"
        ])
        data["Sekundární zdroj"] = select_with_custom("Sekundární zdroj", [
            "krb", "FVE", "není"
        ])

    with st.expander("IV. Inženýrské sítě"):
        data["Voda"] = select_with_custom("Voda", [
            "vodovod", "studna"
        ])
        data["Kanalizace"] = select_with_custom("Kanalizace", [
            "kanalizace", "ČOV", "jímka"
        ])
        data["Dešťová voda"] = select_with_custom("Dešťová voda", [
            "vsakování", "retenční nádrž", "kanalizace"
        ])
        data["Elektřina"] = select_with_custom("Elektřina", [
            "nová přípojka", "stávající"
        ])
        data["Plyn"] = select_with_custom("Plyn", [
            "ano", "ne"
        ])

# 🔹 CHAT HISTORIE
if "messages" not in st.session_state:
    st.session_state.messages = []

for msg in st.session_state.messages:
    with st.chat_message(msg["role"]):
        st.write(msg["content"])

# 🔹 INPUT
user_input = st.chat_input("Napiš dotaz nebo 'vygeneruj zprávu'...")

# 🔥 HLAVNÍ LOGIKA
if user_input:

    wants_doc = any(word in user_input.lower() for word in [
        "zprávu", "zprava", "dokument", "vytvoř", "vygeneruj"
    ])

    if wants_doc:
        errors = []

        if not data.get("Lokalita"):
            errors.append("Vyplň lokalitu")

        if not data.get("Typ území"):
            errors.append("Vyber typ území")

        if not data.get("Počet NP"):
            errors.append("Vyber počet podlaží")

        if errors:
            for e in errors:
                st.error(e)
            st.stop()

    st.session_state.messages.append({"role": "user", "content": user_input})

    # 🔥 RAG jen pro dokument
    if wants_doc:
        docs = vectorstore.similarity_search("technická zpráva stavba projekt", k=5)
        context = "\n\n".join([d.page_content for d in docs])
    else:
        context = ""

    params_text = "\n".join([f"- {k}: {v}" for k, v in data.items() if v])

    # 🔥 PROMPT
    prompt = f"""
Uživatel:
{user_input}

---

Kontext (pokud existuje):
{context}

Parametry projektu:
{params_text}

---

Jsi zkušený projektant a zároveň praktický AI asistent.

---

PRAVIDLA CHOVÁNÍ:

- odpovídej přirozeně, ne jako robot
- jdi rovnou k věci (žádné "rozumím", "rád pomohu")
- vysvětluj jasně a prakticky
- používej strukturované odpovědi (odstavce, odrážky)
- při běžném dotazu můžeš použít i lehké emoji pro přehlednost

---

REŽIMY:

1) BĚŽNÝ DOTAZ (např. "ahoj", "nejde mi iPhone", "co je ČOV"):
- odpověz normálně jako chytrý asistent
- buď konkrétní a užitečný

---

2) TECHNICKÁ ZPRÁVA:

NEJDŘÍV zkontroluj, jestli máš dost informací.

Použij tento checklist:

- Lokalita
- Parcelní číslo
- Typ území
- Počet podlaží
- Půdorys
- Konstrukční systém
- Typ střechy
- Vytápění
- Voda
- Kanalizace
- Elektřina
- Přístup (doprava)

---

POKUD něco chybí:
- NEVYTVOŘ dokument
- polož konkrétní otázky (max 3–4 najednou)

---

POKUD máš vše:
- vytvoř kompletní technickou zprávu
- profesionální styl
- žádné emoji

---

DŮLEŽITÉ:

- nikdy si nevymýšlej chybějící údaje
- raději se zeptej než hádej
- chovej se jako reálný odborník

---
"""

    with st.chat_message("assistant"):
        with st.spinner("Přemýšlím..."):
            response = client.chat.completions.create(
                model="gpt-4o",
                messages=[{"role": "user", "content": prompt}]
            )

            reply = response.choices[0].message.content
            st.write(reply)

            # 🔥 WORD EXPORT
            if wants_doc:
                doc = Document()
                doc.add_heading("TECHNICKÁ ZPRÁVA", 0)

                for line in reply.split("\n"):
                    line = line.strip()
                    if not line:
                        continue

                    if len(line) < 60:
                        doc.add_heading(line, level=1)
                    else:
                        p = doc.add_paragraph(line)
                        p.paragraph_format.space_after = Pt(10)

                filename = "technicka_zprava.docx"
                if data.get("Lokalita"):
                    filename = f"zprava_{data['Lokalita']}.docx"

                doc.save(filename)

                with open(filename, "rb") as f:
                    st.download_button(
                        "📄 Stáhnout jako Word",
                        f,
                        file_name=filename
                    )

            st.session_state.messages.append({
                "role": "assistant",
                "content": reply
            })
